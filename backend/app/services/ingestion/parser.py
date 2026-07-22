"""Unified document parser: converts any supported file into a list of Page objects.

Supported types: PDF, DOCX, TXT, XLSX, PNG/JPG (via OCR fallback for images and
scanned/low-text PDF pages).
"""
from __future__ import annotations

from dataclasses import dataclass, field

from app.core.logging import logger


@dataclass
class ParsedTable:
    page_number: int
    rows: list[list[str]]


@dataclass
class ParsedPage:
    page_number: int
    text: str
    tables: list[ParsedTable] = field(default_factory=list)


@dataclass
class ParsedDocument:
    pages: list[ParsedPage]
    page_count: int
    metadata: dict


class DocumentParser:
    """Routes a file to the correct parsing strategy based on its extension."""

    # A page is considered "text sparse" and routed to OCR if it has fewer
    # than this many extractable characters.
    OCR_TEXT_DENSITY_THRESHOLD = 20

    def parse(self, file_path: str, file_type: str) -> ParsedDocument:
        file_type = file_type.lower().lstrip(".")
        if file_type == "pdf":
            return self._parse_pdf(file_path)
        if file_type == "docx":
            return self._parse_docx(file_path)
        if file_type == "txt":
            return self._parse_txt(file_path)
        if file_type in {"xlsx", "xls"}:
            return self._parse_xlsx(file_path)
        if file_type in {"png", "jpg", "jpeg"}:
            return self._parse_image(file_path)
        raise ValueError(f"Unsupported file type: {file_type}")

    # ------------------------------------------------------------------ PDF
    def _parse_pdf(self, file_path: str) -> ParsedDocument:
        import fitz  # PyMuPDF
        import pdfplumber

        pages: list[ParsedPage] = []

        doc = fitz.open(file_path)
        raw_texts: dict[int, str] = {}
        for i, page in enumerate(doc, start=1):
            raw_texts[i] = page.get_text("text") or ""
        metadata = dict(doc.metadata or {})
        doc.close()

        # Tables via pdfplumber
        tables_by_page: dict[int, list[ParsedTable]] = {}
        try:
            with pdfplumber.open(file_path) as pdf:
                for i, page in enumerate(pdf.pages, start=1):
                    extracted = page.extract_tables() or []
                    if extracted:
                        tables_by_page[i] = [
                            ParsedTable(page_number=i, rows=t) for t in extracted
                        ]
        except Exception as exc:  # pdfplumber can fail on malformed PDFs
            logger.warning(f"pdfplumber table extraction failed for {file_path}: {exc}")

        for page_num, text in raw_texts.items():
            final_text = text
            if len(text.strip()) < self.OCR_TEXT_DENSITY_THRESHOLD:
                # Likely a scanned page — fall back to OCR
                ocr_text = self._ocr_pdf_page(file_path, page_num - 1)
                if ocr_text:
                    final_text = ocr_text
            pages.append(
                ParsedPage(
                    page_number=page_num,
                    text=final_text,
                    tables=tables_by_page.get(page_num, []),
                )
            )

        return ParsedDocument(pages=pages, page_count=len(pages), metadata=metadata)

    def _ocr_pdf_page(self, file_path: str, page_index: int) -> str:
        try:
            import fitz  # PyMuPDF

            from app.services.ingestion.ocr import OCRService

            doc = fitz.open(file_path)
            page = doc[page_index]
            pix = page.get_pixmap(dpi=200)
            image_bytes = pix.tobytes("png")
            doc.close()
            return OCRService().extract_text_from_bytes(image_bytes)
        except Exception as exc:
            logger.warning(f"OCR fallback failed for page {page_index} of {file_path}: {exc}")
            return ""

    # ----------------------------------------------------------------- DOCX
    def _parse_docx(self, file_path: str) -> ParsedDocument:
        import docx

        document = docx.Document(file_path)
        full_text = []
        tables: list[ParsedTable] = []

        for para in document.paragraphs:
            if para.text.strip():
                full_text.append(para.text)

        for table in document.tables:
            rows = [[cell.text for cell in row.cells] for row in table.rows]
            tables.append(ParsedTable(page_number=1, rows=rows))

        # DOCX has no native page concept — treat the whole document as page 1.
        page = ParsedPage(page_number=1, text="\n".join(full_text), tables=tables)
        return ParsedDocument(pages=[page], page_count=1, metadata={})

    # ------------------------------------------------------------------ TXT
    def _parse_txt(self, file_path: str) -> ParsedDocument:
        with open(file_path, encoding="utf-8", errors="ignore") as f:
            text = f.read()
        page = ParsedPage(page_number=1, text=text)
        return ParsedDocument(pages=[page], page_count=1, metadata={})

    # ----------------------------------------------------------------- XLSX
    def _parse_xlsx(self, file_path: str) -> ParsedDocument:
        import pandas as pd

        sheets = pd.read_excel(file_path, sheet_name=None, dtype=str).items()
        pages: list[ParsedPage] = []
        for i, (sheet_name, df) in enumerate(sheets, start=1):
            df = df.fillna("")
            text_repr = f"Sheet: {sheet_name}\n" + df.to_string(index=False)
            rows = [df.columns.tolist()] + df.values.tolist()
            table = ParsedTable(page_number=i, rows=rows)
            pages.append(ParsedPage(page_number=i, text=text_repr, tables=[table]))
        return ParsedDocument(pages=pages, page_count=len(pages), metadata={})

    # ---------------------------------------------------------------- Image
    def _parse_image(self, file_path: str) -> ParsedDocument:
        from app.services.ingestion.ocr import OCRService

        text = OCRService().extract_text_from_file(file_path)
        page = ParsedPage(page_number=1, text=text)
        return ParsedDocument(pages=[page], page_count=1, metadata={})
